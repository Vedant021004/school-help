import os
import re
import uuid
from typing import Dict, Any, List, Optional
import pymupdf
import docx
from backend.models import PaperFormat, SectionFormat


def parse_format_from_text(raw_text: str, filename: str = "Uploaded_Format") -> PaperFormat:
    """
    Parses raw paper format text into a structured PaperFormat model.
    Detects Total Marks, Duration, Sections (Section A, B, C...), Marks per question, and Question types.
    """
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]

    # 1. Total Marks Detection
    total_marks = 50
    marks_match = re.search(r'(total\s+marks|max\s+marks|maximum\s+marks|marks)\s*[:=\-]?\s*(\d+)', raw_text, re.IGNORECASE)
    if marks_match:
        total_marks = int(marks_match.group(2))

    # 2. Duration Detection
    duration_mins = 120
    dur_match = re.search(r'(time|duration)\s*[:=\-]?\s*(\d+(\.\d+)?)\s*(hours?|hrs?|mins?|minutes?)', raw_text, re.IGNORECASE)
    if dur_match:
        val = float(dur_match.group(2))
        unit = dur_match.group(4).lower()
        if "hour" in unit or "hr" in unit:
            duration_mins = int(val * 60)
        else:
            duration_mins = int(val)

    # 3. Section Pattern Detection
    # Look for patterns like "Section A", "Section B", "Part A", "Group A"
    sec_pattern = re.compile(
        r'^(section|part|group)\s+([A-E1-5])[:\-\s]*(.*)$',
        re.IGNORECASE
    )

    # Question line patterns like "10 x 1 = 10 marks", "5 questions of 2 marks each", "MCQ 10 marks"
    calc_pattern = re.compile(r'(\d+)\s*[xX*×]\s*(\d+)\s*=\s*(\d+)', re.IGNORECASE)
    each_pattern = re.compile(r'(\d+)\s*(?:questions?|q)?\s*(?:of|carrying|@)?\s*(\d+)\s*marks?', re.IGNORECASE)

    sections: List[SectionFormat] = []
    current_sec: Optional[Dict[str, Any]] = None
    current_instructions: List[str] = []

    for line in lines:
        m = sec_pattern.match(line)
        if m:
            if current_sec:
                sections.append(SectionFormat(**current_sec))
            sec_letter = m.group(2).upper()
            sec_title = m.group(3).strip() or f"Section {sec_letter}"
            current_sec = {
                "id": f"sec-{sec_letter.lower()}-{uuid.uuid4().hex[:6]}",
                "name": f"Section {sec_letter}",
                "title": sec_title,
                "question_count": 5,
                "marks_per_question": 1,
                "total_marks": 5,
                "question_type": "MCQ" if sec_letter == "A" else "Short Answer",
                "internal_choices_count": 0,
                "instructions": ""
            }
            continue

        # If inside a section, look for question count and marks specifications
        if current_sec:
            # Check for "10 x 1 = 10 marks"
            c_match = calc_pattern.search(line)
            if c_match:
                q_cnt = int(c_match.group(1))
                m_each = int(c_match.group(2))
                current_sec["question_count"] = q_cnt
                current_sec["marks_per_question"] = m_each
                current_sec["total_marks"] = q_cnt * m_each

            # Check for "5 questions of 2 marks each"
            e_match = each_pattern.search(line)
            if e_match and not c_match:
                q_cnt = int(e_match.group(1))
                m_each = int(e_match.group(2))
                current_sec["question_count"] = q_cnt
                current_sec["marks_per_question"] = m_each
                current_sec["total_marks"] = q_cnt * m_each

            # Detect question type
            line_upper = line.upper()
            if "MCQ" in line_upper or "MULTIPLE CHOICE" in line_upper or "OBJECTIVE" in line_upper:
                current_sec["question_type"] = "MCQ"
            elif "VERY SHORT" in line_upper or "VSA" in line_upper:
                current_sec["question_type"] = "Very Short Answer"
            elif "SHORT ANSWER" in line_upper or "SA" in line_upper:
                current_sec["question_type"] = "Short Answer"
            elif "LONG ANSWER" in line_upper or "LA" in line_upper:
                current_sec["question_type"] = "Long Answer"
            elif "CASE STUDY" in line_upper or "SOURCE BASED" in line_upper or "COMPETENCY" in line_upper:
                current_sec["question_type"] = "Case Study"
            elif "NUMERICAL" in line_upper or "PROBLEM" in line_upper:
                current_sec["question_type"] = "Numerical"
            elif "ASSERTION" in line_upper:
                current_sec["question_type"] = "Assertion & Reason"

            if "choice" in line.lower() or "internal" in line.lower() or " or " in line.lower():
                current_sec["internal_choices_count"] = 1

        else:
            # Header instructions
            if any(term in line.lower() for term in ["compulsory", "instructions", "marks indicated", "calculator"]):
                current_instructions.append(line)

    if current_sec:
        sections.append(SectionFormat(**current_sec))

    # Fallback standard sections if no clear sections were parsed
    if not sections:
        sections = [
            SectionFormat(
                name="Section A", title="Multiple Choice Questions",
                question_count=10, marks_per_question=1, total_marks=10, question_type="MCQ"
            ),
            SectionFormat(
                name="Section B", title="Short Answer Type I",
                question_count=5, marks_per_question=2, total_marks=10, question_type="Short Answer"
            ),
            SectionFormat(
                name="Section C", title="Short Answer Type II",
                question_count=5, marks_per_question=3, total_marks=15, question_type="Long Answer"
            ),
            SectionFormat(
                name="Section D", title="Case Study / Problem Solving",
                question_count=3, marks_per_question=5, total_marks=15, question_type="Case Study"
            )
        ]

    # Recalculate total marks if discrepancy
    calc_total = sum(s.total_marks for s in sections)
    if calc_total > 0 and total_marks == 50 and calc_total != 50:
        total_marks = calc_total

    name_clean = os.path.splitext(os.path.basename(filename))[0].replace("_", " ").title()

    return PaperFormat(
        id=f"fmt-{uuid.uuid4().hex[:8]}",
        name=f"Custom Format: {name_clean}",
        description=f"Auto-extracted format with {len(sections)} sections and {total_marks} total marks.",
        total_marks=total_marks,
        duration_minutes=duration_mins,
        instructions=current_instructions if current_instructions else [
            "All questions are compulsory.",
            "Marks are indicated against each question.",
            "Write answers neatly and legibly."
        ],
        sections=sections,
        is_template=False
    )


def extract_format_from_file(file_path: str) -> PaperFormat:
    """
    Extracts text from PDF, DOCX, or text file and parses paper format.
    """
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    if ext == ".pdf":
        doc = pymupdf.open(file_path)
        for page in doc:
            extracted_text += page.get_text("text") + "\n"
        doc.close()
    elif ext in [".docx", ".doc"]:
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            extracted_text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                extracted_text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            extracted_text = f.read()

    return parse_format_from_text(extracted_text, filename=file_path)
